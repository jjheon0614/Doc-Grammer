""" Filename: docx_utils.py - Directory: my_flask_app/utils
"""
import os
import asyncio
import nltk
from flask import current_app
from docx import Document
from aiohttp import ClientSession
from .grammar_checker import check_grammar
from .reconstructing_sentence import *

# Download necessary Punkt sentence tokenizer if not already downloaded
nltk.download("punkt", quiet=True)


async def correct_text_grammar(file_path):
    doc = Document(file_path)
    corrections = []  # List to hold correction details
    corrected_paragraphs = (
        {}
    )  # Use a dictionary to map paragraph index to corrected text
    excluded_sections = [
        "endnote bibliography",
        "icce affiliations",
        "icce author list",
    ]

    async with ClientSession() as session:
        for i in range(0, len(doc.paragraphs), 10):
            tasks = []
            batch_indices = []  # Store indices of paragraphs in the current batch

            for index, paragraph in enumerate(doc.paragraphs[i : i + 10], start=i):
                style_name = paragraph.style.name.lower()
                paragraph_text = paragraph.text.strip()

                if not paragraph_text:
                    continue  # Skip empty paragraphs

                if is_code_snippet(paragraph_text):
                    continue  # Skip code snippets

                if len(paragraph_text) <= 40:
                    continue  # Skip paragraphs that are too short to be meaningful

                if style_name in excluded_sections:
                    continue  # Skip specific sections

                task = asyncio.create_task(process_paragraph(paragraph, session))
                tasks.append(task)
                batch_indices.append(
                    index
                )  # Keep track of the paragraph's original index

            # Wait for all tasks in the current batch to complete and store results
            results = await asyncio.gather(*tasks)

            for index, corrected_text in zip(batch_indices, results):
                original_text = doc.paragraphs[index].text
                # Check if the text was modified and if the modification is more than just an extra period at the end
                if original_text != corrected_text and not (
                    original_text + "." == corrected_text
                    or corrected_text + "." == original_text
                    or corrected_text == ""
                ):
                    corrections.append(
                        {
                            "original_sentence": original_text,
                            "corrected_sentence": corrected_text,
                        }
                    )

                corrected_paragraphs[
                    index
                ] = corrected_text  # Map index to corrected text

            # print("Processed batch of paragraphs. Waiting before next batch...")
            await asyncio.sleep(0.1)

        # Apply the corrected text to each paragraph in the original order
        for index, paragraph in enumerate(doc.paragraphs):
            if (
                paragraph.text.strip()
                and paragraph.style.name != "EndNote Bibliography"
                and index in corrected_paragraphs
            ):
                # paragraph.text = corrected_paragraphs[index]
                correct_paragraph(corrected_paragraphs[index], paragraph)

    # Save the corrected document
    doc.save(file_path)
    print("Completed grammar correction and saved document.")
    # Return all corrections details
    return corrections


async def process_paragraph(paragraph, session):
    corrected_para = ""
    para_word_count = len(paragraph.text.split())

    # Function to check for hyperlinks
    def contains_hyperlink(text):
        return re.search(r"https?://\S+", text) is not None

    if para_word_count <= 64:
        # Check for hyperlink in short paragraphs
        if not contains_hyperlink(paragraph.text):
            corrected_para = await async_check_grammar(paragraph.text, session)
            await asyncio.sleep(0.01)  # Delay between batches

    else:
        # If the paragraph is long, split into sentences
        sentences = nltk.tokenize.sent_tokenize(paragraph.text)
        sentence_batches = [sentences[i : i + 10] for i in range(0, len(sentences), 10)]

        for batch in sentence_batches:
            processed_sentences = []
            for sentence in batch:
                if not contains_hyperlink(sentence):
                    processed_sentence = await async_check_grammar(sentence, session)
                    processed_sentences.append(processed_sentence)
                else:
                    processed_sentences.append(
                        sentence
                    )  # Skip processing, keep original

            corrected_para += " ".join(processed_sentences)
            await asyncio.sleep(0.01)  # Delay between batches

    return corrected_para


def is_code_snippet(paragraph_text):
    return is_java_code_snippet(paragraph_text) or is_html_code_snippet(paragraph_text)


def is_java_code_snippet(paragraph_text):
    # Simple checks for code-like structures
    code_indicators = [
        r"\s*;\s*$",  # lines ending in a semicolon with optional leading whitespace
        r"\s*{\s*$",  # line ending with an opening curly brace with optional leading whitespace
        r"\s*}\s*$",  # line ending with a closing curly brace with optional leading whitespace
        r"\s*class\s+",  # lines that start with 'class' keyword
        r"\s*public\s+",  # lines that start with 'public' keyword
        r"\s*private\s+",  # lines that start with 'private' keyword
        r"\s*(if|for|while|switch|return)\b",  # common control structures
    ]
    for indicator in code_indicators:
        if re.search(indicator, paragraph_text):
            return True
    return False


def is_html_code_snippet(paragraph_text):
    # Simple checks for HTML code-like structures
    html_code_indicators = [
        r"\s*<\s*/?\s*[a-zA-Z][^>]*>",  # Detects HTML tags with optional leading whitespace
        r"\s*<\s*[a-zA-Z]+[^>]*>\s*</\s*[a-zA-Z]+\s*>",  # Detects HTML open and close tag pairs with optional leading whitespace
        r"\s*<!DOCTYPE\s+html>",  # Detects DOCTYPE declaration with optional leading whitespace
        r"\s*<!--.*?-->",  # Detects HTML comments with optional leading whitespace
    ]
    for indicator in html_code_indicators:
        if re.search(indicator, paragraph_text, re.DOTALL | re.MULTILINE):
            return True
    return False


async def async_check_grammar(original_text, session):
    corrected_text = await check_grammar(original_text, session)
    return corrected_text


def correct_paragraph(corrected_para, paragraph):
    modifications = find_modified_text(paragraph.text, corrected_para)
    paragraph_tokens = custom_tokenize(paragraph.text)

    modifications_dict = {item[2]: (item[0], item[1]) for item in modifications}
    # print("\nModifications Dictionary:", modifications_dict)

    current_token_index = 0
    current_pos = 0

    for run_index, run in enumerate(paragraph.runs):
        # print(f"\nProcessing run {run_index} (Text: '{run.text}')")
        new_run_text = ""
        run_start_pos = current_pos
        run_end_pos = current_pos + len(run.text)

        while current_token_index < len(paragraph_tokens) and current_pos < run_end_pos:
            token = paragraph_tokens[current_token_index]
            token_end_pos = current_pos + len(token)
            # print(
            #     f"  Current token: '{token}' (Index: {current_token_index}, Positions: {current_pos}-{token_end_pos})"
            # )

            if current_token_index in modifications_dict:
                original, modified = modifications_dict[current_token_index]
                # print(f"    Modification found: '{original}' -> '{modified}'")

                if token_end_pos <= run_end_pos:
                    # print(f"    Applying full modification to run {run_index}")
                    new_run_text += modified
                    current_token_index += 1
                    current_pos = token_end_pos
                elif run_start_pos < token_end_pos and current_pos < run_end_pos:
                    partial_mod_length = run_end_pos - current_pos
                    # print(
                    #     f"    Applying partial modification to run {run_index}: '{modified[:partial_mod_length]}'"
                    # )
                    new_run_text += modified[:partial_mod_length]

                    if len(original) > partial_mod_length:
                        modifications_dict[current_token_index] = (
                            original[partial_mod_length:],
                            modified[partial_mod_length:],
                        )
                    else:
                        # print("    Skipping token, not in current run")
                        current_token_index += 1

                    current_pos = run_end_pos
                else:
                    break
            else:
                if run_start_pos <= current_pos < run_end_pos:
                    # print(f"    Adding unmodified token: '{token}'")
                    new_run_text += token
                    current_token_index += 1
                    current_pos = token_end_pos
                else:
                    # print("    Token not in current run, moving to next run")
                    break

        run.text = new_run_text
        current_pos = run_end_pos
        # print(f"  Updated run {run_index} text: '{run.text}'")


def find_modified_text(original_text, corrected_text):
    diff = percentage_difference(original_text, corrected_text)

    # Check if the text difference is more than 45%
    if diff > 45:
        return []

    mod = []
    ori_tokens = custom_tokenize(original_text)
    cor_tokens = custom_tokenize(corrected_text)
    ori_matrix = [[token, index] for index, token in enumerate(ori_tokens)]
    cor_matrix = [[token, index] for index, token in enumerate(cor_tokens)]

    temp_mod = find_modified_tokens(ori_matrix, cor_matrix)

    # TODO: replace ori_tokens with mod and then rerun the find_modified_tokens again to improve correctness....
    for item in temp_mod:
        index = item[2]
        ori_tokens[index] = item[1]

    updated_ori_matrix = [[token, index] for index, token in enumerate(ori_tokens)]

    mod = find_modified_tokens(updated_ori_matrix, cor_matrix)
    combined_list = temp_mod + mod

    final_mod = sorted(list(set(combined_list)), key=lambda x: x[2])
    final_mod = [item for item in final_mod if item[0] != item[1]]

    return final_mod


def find_modified_tokens(ori, cor):
    ori_index = cor_index = 0
    mod = []

    while ori_index < len(ori) and cor_index < len(cor):
        # print(f"Current Iteration: {ori_index}, {cor_index}")
        ori_token, ori_pos = ori[ori_index]
        cor_token, _ = cor[cor_index]

        if ori_token == cor_token:
            # Tokens are identical, move to the next pair
            ori_index += 1
            cor_index += 1
        else:
            if ori_index + 4 < len(ori) and cor_index + 4 < len(cor):
                if ori[ori_index + 4][0] == cor[cor_index + 2][0]:
                    mod.append((ori_token, "", ori_pos))
                    mod.append((" ", "", ori_pos + 1))
                    mod.append((ori[ori_index + 2][0], cor_token, ori_pos + 2))

                    ori_index += 3
                    cor_index += 1
                    continue
                elif ori[ori_index + 2][0] == cor[cor_index + 2][0]:
                    mod.append((ori_token, cor_token, ori_pos))
                    ori_index += 1
                    cor_index += 1
                    continue
                elif ori[ori_index + 2][0] == cor[cor_index + 4][0]:
                    mod.append(
                        (ori_token, cor_token + " " + cor[cor_index + 2][0], ori_pos)
                    )
                    ori_index += 1
                    cor_index += 3
                    continue
                elif (
                    ori[ori_index + 2][0] != cor[cor_index + 2][0]
                    and ori[ori_index + 4][0] == cor[cor_index + 4][0]
                ):
                    mod.append((ori_token, cor_token, ori_pos))
                    mod.append(
                        (ori[ori_index + 2][0], cor[cor_index + 2][0], ori_pos + 2)
                    )
                    ori_index += 3
                    cor_index += 3

                else:
                    ori_index += 1
                    cor_index += 1
            else:
                # Handle end-of-text cases
                if len(ori) - ori_index == len(cor) - cor_index:
                    mod.append((ori_token, cor_token, ori_pos))
                    ori_index += 1
                    cor_index += 1
                elif len(ori) - ori_index > len(cor) - cor_index:
                    mod.append((ori_token, "", ori_pos))
                    mod.append((" ", "", ori_pos + 1))
                    mod.append((ori[ori_index + 2][0], cor_token, ori_pos + 2))
                    ori_index += 3
                    cor_index += 1
                else:
                    if cor_index + 2 < len(cor):
                        mod.append(
                            (
                                ori_token,
                                cor_token + " " + cor[cor_index + 2][0],
                                ori_pos,
                            )
                        )
                        cor_index += 3
                    else:
                        mod.append((ori_token, cor_token, ori_pos))
                        cor_index += 1
                    ori_index += 1

    return mod
